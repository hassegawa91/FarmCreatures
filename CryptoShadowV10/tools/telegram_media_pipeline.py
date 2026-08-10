from __future__ import annotations

import argparse
import json
import os
import re
import sqlite3
import subprocess
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable


IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".gif"}
AUDIO_EXTENSIONS = {".ogg", ".mp3", ".wav", ".m4a"}
VIDEO_EXTENSIONS = {".mp4", ".mov", ".mkv", ".webm"}
DOCUMENT_EXTENSIONS = {".pdf", ".txt", ".docx", ".xlsx", ".csv"}


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def initialize(connection: sqlite3.Connection) -> None:
    connection.executescript(
        """
        PRAGMA journal_mode=WAL;
        PRAGMA synchronous=NORMAL;
        CREATE TABLE IF NOT EXISTS media_analysis (
            sha256 TEXT NOT NULL,
            stage TEXT NOT NULL,
            checkpoint_id TEXT NOT NULL,
            source_path TEXT NOT NULL,
            extension TEXT NOT NULL,
            status TEXT NOT NULL,
            extracted_text TEXT NOT NULL DEFAULT '',
            metadata_json TEXT NOT NULL DEFAULT '{}',
            error TEXT,
            processed_utc TEXT NOT NULL,
            elapsed_seconds REAL NOT NULL DEFAULT 0,
            PRIMARY KEY (sha256, stage)
        );
        CREATE INDEX IF NOT EXISTS media_analysis_stage ON media_analysis(stage, status);
        CREATE VIRTUAL TABLE IF NOT EXISTS media_fts USING fts5(
            sha256 UNINDEXED, stage UNINDEXED, source_path, extracted_text,
            tokenize='unicode61 remove_diacritics 2'
        );
        """
    )


def source_root(connection: sqlite3.Connection, checkpoint_id: str) -> Path:
    row = connection.execute(
        "SELECT source_root FROM checkpoints WHERE checkpoint_id=?", (checkpoint_id,)
    ).fetchone()
    if not row:
        raise RuntimeError(f"Checkpoint desconhecido: {checkpoint_id}")
    return Path(row[0])


def media_artifacts(connection: sqlite3.Connection, checkpoint_id: str) -> list[dict[str, Any]]:
    rows = connection.execute(
        """SELECT sha256, MIN(path), MAX(bytes)
           FROM artifacts WHERE checkpoint_id=? AND kind='MEDIA'
           GROUP BY sha256 ORDER BY MIN(path)""",
        (checkpoint_id,),
    ).fetchall()
    return [{"sha256": row[0], "path": row[1], "bytes": row[2]} for row in rows]


def already_processed(connection: sqlite3.Connection, sha256: str, stage: str) -> bool:
    return connection.execute(
        "SELECT 1 FROM processed_hashes WHERE sha256=? AND stage=?", (sha256, stage)
    ).fetchone() is not None


def append_ledger(ledger_path: Path, entry: dict[str, Any]) -> None:
    ledger_path.parent.mkdir(parents=True, exist_ok=True)
    with ledger_path.open("a", encoding="utf-8") as ledger:
        ledger.write(json.dumps(entry, ensure_ascii=False) + "\n")


def save_result(
    connection: sqlite3.Connection,
    ledger_path: Path,
    *,
    checkpoint_id: str,
    sha256: str,
    stage: str,
    path: str,
    status: str,
    text: str = "",
    metadata: dict[str, Any] | None = None,
    error: str | None = None,
    elapsed: float = 0.0,
) -> None:
    processed = utc_now()
    extension = Path(path).suffix.lower()
    metadata_json = json.dumps(metadata or {}, ensure_ascii=False)
    connection.execute(
        """INSERT OR REPLACE INTO media_analysis(
            sha256,stage,checkpoint_id,source_path,extension,status,extracted_text,
            metadata_json,error,processed_utc,elapsed_seconds
        ) VALUES(?,?,?,?,?,?,?,?,?,?,?)""",
        (sha256, stage, checkpoint_id, path, extension, status, text, metadata_json, error, processed, elapsed),
    )
    if text:
        connection.execute(
            "INSERT INTO media_fts(sha256,stage,source_path,extracted_text) VALUES(?,?,?,?)",
            (sha256, stage, path, text),
        )
    connection.execute(
        "INSERT OR REPLACE INTO processed_hashes VALUES(?,?,?,?,?,?)",
        (sha256, stage, processed, checkpoint_id, path, 1),
    )
    connection.execute(
        "UPDATE artifacts SET status=? WHERE checkpoint_id=? AND sha256=?",
        (stage if status == "OK" else status, checkpoint_id, sha256),
    )
    connection.commit()
    append_ledger(ledger_path, {
        "sha256": sha256, "stage": stage, "processed_utc": processed,
        "checkpoint_id": checkpoint_id, "path": path, "item_count": 1,
        "status": status, "error": error,
    })


def image_metadata(path: Path) -> dict[str, Any]:
    from PIL import Image
    import imagehash

    with Image.open(path) as image:
        return {
            "width": image.width,
            "height": image.height,
            "mode": image.mode,
            "format": image.format,
            "frames": getattr(image, "n_frames", 1),
            "perceptual_hash": str(imagehash.phash(image.convert("RGB"))),
        }


def av_metadata(path: Path) -> dict[str, Any]:
    import av

    with av.open(str(path)) as container:
        streams = []
        for stream in container.streams:
            item: dict[str, Any] = {
                "type": stream.type,
                "codec": stream.codec_context.name,
                "duration_seconds": float(stream.duration * stream.time_base) if stream.duration else None,
            }
            if stream.type == "video":
                item.update({"width": stream.codec_context.width, "height": stream.codec_context.height})
            if stream.type == "audio":
                item.update({"sample_rate": stream.codec_context.sample_rate, "channels": stream.codec_context.channels})
            streams.append(item)
        duration = float(container.duration / 1_000_000) if container.duration else None
        return {"duration_seconds": duration, "streams": streams, "format": container.format.name}


def pdf_metadata(path: Path) -> dict[str, Any]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    return {
        "pages": len(reader.pages),
        "encrypted": reader.is_encrypted,
        "metadata": {str(key): str(value) for key, value in (reader.metadata or {}).items()},
    }


def run_metadata(
    connection: sqlite3.Connection, checkpoint_id: str, root: Path, ledger: Path,
    artifacts: Iterable[dict[str, Any]], limit: int | None,
) -> dict[str, int]:
    stage = "METADATA"
    done = skipped = failed = 0
    for artifact in artifacts:
        if limit is not None and done + failed >= limit:
            break
        sha = artifact["sha256"]
        if already_processed(connection, sha, stage):
            skipped += 1
            continue
        relative = artifact["path"]
        path = root / relative
        suffix = path.suffix.lower()
        started = time.perf_counter()
        try:
            if suffix in IMAGE_EXTENSIONS:
                metadata = image_metadata(path)
            elif suffix in AUDIO_EXTENSIONS | VIDEO_EXTENSIONS:
                metadata = av_metadata(path)
            elif suffix == ".pdf":
                metadata = pdf_metadata(path)
            elif suffix == ".txt":
                metadata = {"encoding": "utf-8-or-fallback", "bytes": path.stat().st_size}
            else:
                metadata = {"bytes": path.stat().st_size}
            save_result(connection, ledger, checkpoint_id=checkpoint_id, sha256=sha, stage=stage,
                        path=relative, status="OK", metadata=metadata,
                        elapsed=time.perf_counter() - started)
            done += 1
        except Exception as exc:
            save_result(connection, ledger, checkpoint_id=checkpoint_id, sha256=sha, stage=stage,
                        path=relative, status="ERROR", error=f"{type(exc).__name__}: {exc}",
                        elapsed=time.perf_counter() - started)
            failed += 1
    return {"processed": done, "skipped": skipped, "failed": failed}


def extract_pdf_text(path: Path) -> tuple[str, dict[str, Any]]:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(path) as document:
        for number, page in enumerate(document.pages, start=1):
            text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
            pages.append(f"\n\n===== PAGE {number} =====\n{text}")
        return "".join(pages).strip(), {"pages": len(document.pages), "page_characters": [len(x) for x in pages]}


def extract_docx_text(path: Path) -> tuple[str, dict[str, Any]]:
    from docx import Document

    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables: list[str] = []
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                tables.append(" | ".join(values))
    text = "\n".join(paragraphs + (["\n===== TABLES ====="] + tables if tables else []))
    return text, {"paragraphs": len(paragraphs), "table_rows": len(tables)}


def extract_xlsx_text(path: Path) -> tuple[str, dict[str, Any]]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    lines: list[str] = []
    row_count = 0
    for sheet in workbook.worksheets:
        lines.append(f"\n===== SHEET {sheet.title} =====")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                lines.append(" | ".join(values))
                row_count += 1
    workbook.close()
    return "\n".join(lines).strip(), {"sheets": len(workbook.sheetnames), "rows": row_count}


def run_documents(
    connection: sqlite3.Connection, checkpoint_id: str, root: Path, ledger: Path,
    artifacts: Iterable[dict[str, Any]], limit: int | None,
) -> dict[str, int]:
    stage = "DOCUMENT_TEXT"
    done = skipped = failed = 0
    for artifact in artifacts:
        relative = artifact["path"]
        suffix = Path(relative).suffix.lower()
        if suffix not in DOCUMENT_EXTENSIONS:
            continue
        if limit is not None and done + failed >= limit:
            break
        sha = artifact["sha256"]
        if already_processed(connection, sha, stage):
            skipped += 1
            continue
        path = root / relative
        started = time.perf_counter()
        try:
            if suffix == ".pdf":
                text, metadata = extract_pdf_text(path)
            elif suffix == ".docx":
                text, metadata = extract_docx_text(path)
            elif suffix == ".xlsx":
                text, metadata = extract_xlsx_text(path)
            else:
                raw = path.read_bytes()
                text = raw.decode("utf-8", errors="replace")
                metadata = {"characters": len(text)}
            save_result(connection, ledger, checkpoint_id=checkpoint_id, sha256=sha, stage=stage,
                        path=relative, status="OK", text=text, metadata=metadata,
                        elapsed=time.perf_counter() - started)
            done += 1
        except Exception as exc:
            save_result(connection, ledger, checkpoint_id=checkpoint_id, sha256=sha, stage=stage,
                        path=relative, status="ERROR", error=f"{type(exc).__name__}: {exc}",
                        elapsed=time.perf_counter() - started)
            failed += 1
    return {"processed": done, "skipped": skipped, "failed": failed}


def load_image(path: Path):
    import cv2
    import numpy as np

    return cv2.imdecode(np.fromfile(path, dtype=np.uint8), cv2.IMREAD_COLOR)


def ocr_array(engine, image) -> tuple[str, dict[str, Any]]:
    result, elapsed = engine(image)
    if not result:
        return "", {"lines": 0, "ocr_seconds": elapsed}
    lines = [str(item[1]) for item in result if len(item) >= 2]
    scores = [float(item[2]) for item in result if len(item) >= 3]
    return "\n".join(lines), {
        "lines": len(lines), "mean_confidence": sum(scores) / len(scores) if scores else None,
        "ocr_seconds": elapsed,
    }


def run_image_ocr(
    connection: sqlite3.Connection, checkpoint_id: str, root: Path, ledger: Path,
    artifacts: Iterable[dict[str, Any]], limit: int | None,
) -> dict[str, int]:
    from rapidocr_onnxruntime import RapidOCR

    stage = "IMAGE_OCR"
    # Keep two logical CPUs free for the live trading engine and dashboard.
    engine = RapidOCR(intra_op_num_threads=2, inter_op_num_threads=1)
    done = skipped = failed = derived = 0
    for artifact in artifacts:
        relative = artifact["path"]
        if Path(relative).suffix.lower() not in IMAGE_EXTENSIONS:
            continue
        if limit is not None and done + failed + derived >= limit:
            break
        sha = artifact["sha256"]
        if already_processed(connection, sha, stage):
            skipped += 1
            continue
        started = time.perf_counter()
        if "_thumb." in relative.lower():
            original = relative.replace("_thumb.", ".")
            save_result(connection, ledger, checkpoint_id=checkpoint_id, sha256=sha, stage=stage,
                        path=relative, status="DERIVED_THUMBNAIL", metadata={"original_path": original},
                        elapsed=time.perf_counter() - started)
            derived += 1
            continue
        try:
            image = load_image(root / relative)
            if image is None:
                raise ValueError("imagem nao decodificada")
            text, metadata = ocr_array(engine, image)
            metadata.update({"width": int(image.shape[1]), "height": int(image.shape[0])})
            save_result(connection, ledger, checkpoint_id=checkpoint_id, sha256=sha, stage=stage,
                        path=relative, status="OK", text=text, metadata=metadata,
                        elapsed=time.perf_counter() - started)
            done += 1
        except Exception as exc:
            save_result(connection, ledger, checkpoint_id=checkpoint_id, sha256=sha, stage=stage,
                        path=relative, status="ERROR", error=f"{type(exc).__name__}: {exc}",
                        elapsed=time.perf_counter() - started)
            failed += 1
    return {"processed": done, "derived_thumbnails": derived, "skipped": skipped, "failed": failed}


def run_transcription(
    connection: sqlite3.Connection, checkpoint_id: str, root: Path, ledger: Path,
    artifacts: Iterable[dict[str, Any]], limit: int | None, model_name: str,
    beam_size: int = 1, cpu_threads: int = 3,
) -> dict[str, int]:
    from faster_whisper import WhisperModel

    stage = f"TRANSCRIPT_{model_name.upper().replace('-', '_')}"
    model = WhisperModel(
        model_name, device="cpu", compute_type="int8",
        cpu_threads=max(1, cpu_threads), num_workers=1,
    )
    metadata_rows = connection.execute(
        "SELECT sha256,metadata_json FROM media_analysis WHERE stage='METADATA' AND status='OK'"
    ).fetchall()
    durations: dict[str, float] = {}
    for sha256, raw in metadata_rows:
        try:
            durations[sha256] = float(json.loads(raw).get("duration_seconds") or 0.0)
        except (TypeError, ValueError, json.JSONDecodeError):
            durations[sha256] = 0.0
    strategic = re.compile(
        r"(?:aula|setup|encryptos|gerenc|risco|trade|tend|romp|revers|"
        r"open.?interest|\boi\b|lsr|volume|liquid|fibo|candle|vela)", re.IGNORECASE,
    )
    candidates = [
        artifact for artifact in artifacts
        if Path(artifact["path"]).suffix.lower() in AUDIO_EXTENSIONS | VIDEO_EXTENSIONS
    ]
    candidates.sort(key=lambda artifact: (
        0 if strategic.search(artifact["path"]) else 1,
        durations.get(artifact["sha256"], float("inf")) or float("inf"),
        artifact["path"].lower(),
    ))
    done = skipped = failed = 0
    for artifact in candidates:
        relative = artifact["path"]
        if limit is not None and done + failed >= limit:
            break
        sha = artifact["sha256"]
        if already_processed(connection, sha, stage):
            skipped += 1
            continue
        started = time.perf_counter()
        try:
            segments, info = model.transcribe(
                str(root / relative), language="pt", vad_filter=True,
                beam_size=max(1, beam_size), best_of=1,
                condition_on_previous_text=False,
            )
            collected = []
            segment_rows = []
            for segment in segments:
                text = segment.text.strip()
                if text:
                    collected.append(text)
                    segment_rows.append({"start": segment.start, "end": segment.end, "text": text})
            transcript = " ".join(collected)
            metadata = {
                "language": info.language, "language_probability": info.language_probability,
                "duration_seconds": info.duration, "segments": segment_rows,
                "model": model_name, "beam_size": beam_size, "cpu_threads": cpu_threads,
            }
            save_result(connection, ledger, checkpoint_id=checkpoint_id, sha256=sha, stage=stage,
                        path=relative, status="OK", text=transcript, metadata=metadata,
                        elapsed=time.perf_counter() - started)
            done += 1
        except Exception as exc:
            save_result(connection, ledger, checkpoint_id=checkpoint_id, sha256=sha, stage=stage,
                        path=relative, status="ERROR", error=f"{type(exc).__name__}: {exc}",
                        elapsed=time.perf_counter() - started)
            failed += 1
    return {"processed": done, "skipped": skipped, "failed": failed}


def run_video_frame_ocr(
    connection: sqlite3.Connection, checkpoint_id: str, root: Path, ledger: Path,
    artifacts: Iterable[dict[str, Any]], limit: int | None, interval_seconds: float,
) -> dict[str, int]:
    import cv2
    from rapidocr_onnxruntime import RapidOCR

    stage = "VIDEO_FRAME_OCR"
    # Video-frame OCR follows the same CPU budget as still-image OCR.
    engine = RapidOCR(intra_op_num_threads=2, inter_op_num_threads=1)
    done = skipped = failed = 0
    for artifact in artifacts:
        relative = artifact["path"]
        if Path(relative).suffix.lower() not in VIDEO_EXTENSIONS:
            continue
        if limit is not None and done + failed >= limit:
            break
        sha = artifact["sha256"]
        if already_processed(connection, sha, stage):
            skipped += 1
            continue
        started = time.perf_counter()
        capture = cv2.VideoCapture(str(root / relative))
        try:
            fps = capture.get(cv2.CAP_PROP_FPS) or 25.0
            frames = int(capture.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
            duration = frames / fps if frames else 0.0
            moments = [0.0]
            position = interval_seconds
            while position < duration:
                moments.append(position)
                position += interval_seconds
            texts: list[str] = []
            sampled: list[dict[str, Any]] = []
            for moment in moments:
                capture.set(cv2.CAP_PROP_POS_MSEC, moment * 1000.0)
                ok, frame = capture.read()
                if not ok:
                    continue
                # Telegram screen recordings can be 4K. OCR at native size is
                # needlessly expensive and previously saturated the machine.
                max_width = 1600
                if frame.shape[1] > max_width:
                    scale = max_width / frame.shape[1]
                    frame = cv2.resize(
                        frame, (max_width, max(1, int(frame.shape[0] * scale))),
                        interpolation=cv2.INTER_AREA,
                    )
                text, metadata = ocr_array(engine, frame)
                if text:
                    texts.append(f"[FRAME {moment:.1f}s]\n{text}")
                sampled.append({"second": moment, **metadata})
            save_result(connection, ledger, checkpoint_id=checkpoint_id, sha256=sha, stage=stage,
                        path=relative, status="OK", text="\n\n".join(texts),
                        metadata={"duration_seconds": duration, "sample_interval_seconds": interval_seconds,
                                  "frames_sampled": sampled},
                        elapsed=time.perf_counter() - started)
            done += 1
        except Exception as exc:
            save_result(connection, ledger, checkpoint_id=checkpoint_id, sha256=sha, stage=stage,
                        path=relative, status="ERROR", error=f"{type(exc).__name__}: {exc}",
                        elapsed=time.perf_counter() - started)
            failed += 1
        finally:
            capture.release()
    return {"processed": done, "skipped": skipped, "failed": failed}


def main() -> None:
    parser = argparse.ArgumentParser(description="Pipeline incremental de midias do Telegram")
    parser.add_argument("database", type=Path)
    parser.add_argument("ledger", type=Path)
    parser.add_argument("checkpoint_id")
    parser.add_argument("stage", choices=["metadata", "documents", "images", "transcribe", "video-frames"])
    parser.add_argument("--limit", type=int)
    parser.add_argument("--model", default="small")
    parser.add_argument("--beam-size", type=int, default=1)
    parser.add_argument("--cpu-threads", type=int, default=max(1, (os.cpu_count() or 2) - 1))
    parser.add_argument("--frame-interval", type=float, default=15.0)
    args = parser.parse_args()

    connection = sqlite3.connect(args.database)
    initialize(connection)
    root = source_root(connection, args.checkpoint_id)
    artifacts = media_artifacts(connection, args.checkpoint_id)
    if args.stage == "metadata":
        result = run_metadata(connection, args.checkpoint_id, root, args.ledger, artifacts, args.limit)
    elif args.stage == "documents":
        result = run_documents(connection, args.checkpoint_id, root, args.ledger, artifacts, args.limit)
    elif args.stage == "images":
        result = run_image_ocr(connection, args.checkpoint_id, root, args.ledger, artifacts, args.limit)
    elif args.stage == "transcribe":
        result = run_transcription(
            connection, args.checkpoint_id, root, args.ledger, artifacts, args.limit,
            args.model, args.beam_size, args.cpu_threads,
        )
    else:
        result = run_video_frame_ocr(connection, args.checkpoint_id, root, args.ledger, artifacts, args.limit, args.frame_interval)
    connection.close()
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
